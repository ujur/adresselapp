# encoding: UTF-8
import os
import sys
import time
import subprocess
print('Python' + sys.version)
# Install dependencies if required
try:
    from ldap3 import Server, Connection, ALL
    from lxml import html
    import requests
    from docx import Document
    from docx.shared import Pt
except ImportError as e:
    print(e)
    # subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-r', 'requirements.txt'])
    subprocess.run([sys.executable, '-m', 'pip', 'install', '--user',
                    'ldap3',
                    'requests',
                    'lxml',
                    'python-docx==0.8.11'
                    ],
                   check=True)
    print("Software installed, restart program. Exiting in 5 seconds.")
    time.sleep(5)
    exit(0)


def get_input(prompt):
    "Get a nonempty input from the user"
    read = ""
    while not read:
        read = input(prompt).strip()
    return read


def print_word_file(filename):
    "Send a word file to the default Windows printer"
    try:
        print("printing...")
        subprocess.run(["write", "/p", filename], check=True)
#         subprocess.run(["write", filename])
    except Exception as e:
        print("Unable to print document:", e)


def get_user_ou(user):
    "Get the printable organisation unit of user"
    return format_ou(user.eduPersonPrimaryOrgUnitDN)


def xstr(item):
    """
    Get string representation, or empty string if item is None
    Example:
    >>> xstr(None)
    ''
    >>> xstr("test")
    'test'
    """
    if item:
        return str(item)
    else:
        return ""


def make_criteria(string):
    """
    Make conjunctive search criteria from a name string

    Example:
    >>> make_criteria("John ")
    '(&(cn=*John*))'
    >>> make_criteria(" John Doe Henry ")
    '(&(cn=*John*)(cn=*Doe*)(cn=*Henry*))'
    """
    terms = ["(cn=*%s*)" % name.strip(",.") for name in string.split()if name]
    return "(&%s)" % "".join(terms)


def format_ou(ou_string):
    """
    Make a human-readable string of the OU

    Example:
    >>> format_ou("ou=UJUR,ou=UB,ou=UIO,cn=organization,dc=uio,dc=no")
    'UJUR/UB/UIO'
    """
    parts = [item[3:] for item in xstr(ou_string).split(",") if item.startswith("ou=")]
    return "/".join(parts)


def get_address_from_web(username):
    """
    Get the correct address from the person profile webpage
    Parameters:
        username of the person
    Example:
    >>> get_address_from_web("ewinge")
    ['Karl Johans gate 47', 'Domus Bibliotheca']
    >>> get_address_from_web("invalid-user-name")
    []
    """
    URL = "http://www.uio.no/?vrtx=person-view&uid=" + username
    page = requests.get(URL)
    tree = html.fromstring(page.content)
    address = tree.xpath(
        '//div[contains(@class,"vrtx-person-visiting-address")]/span[@class="vrtx-address-line"]/text()')
    # remove postcode
    if address:
        address = address[:-1] if len(address) > 1 else [address[0].strip()]
    return address


def print_person(entry):
    "print address slip for an LDAP entry"
    filename = os.path.abspath("address-temp.docx")

    address = [xstr(entry.cn)]
    address.append(get_user_ou(entry))
    address.append("\n".join(get_address_from_web(str(entry.uid))))
    address = "\n".join(address)
    print(address)
    # Create and print docx document
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(address)
#     run.font.name = 'Calibri'
    run.font.size = Pt(16)
    try:
        document.save(filename)
        print_word_file(filename)
        os.remove(filename)
    except Exception as e:
        print(e)


def prompt():
    "Display prompt to user"
    name = get_input("Navn: ").strip()
    if name in["quit", "exit"]:
        exit(0)
    find_person(name)


def find_person(name):
    "Lookup a user in LDAP"
    query = "(cn=*" + name + "*)"
#     print(query)
    with(Connection("ldap.uio.no", auto_bind=True)) as con:
        con.search(
            "cn=people,dc=uio,dc=no",
            query,
            attributes=[
                "uid",
                "cn",
                "postalAddress",
                "eduPersonPrimaryOrgUnitDN",
                "street",
                "uioShortPhone"])
#         print(con.entries)

        for index, user in enumerate(con.entries):
            print(
                "[%2d] %-35s %-25s %s" %
                (index,
                 user.cn,
                 get_user_ou(user),
                 xstr(
                     user.uioShortPhone)))

        if len(con.entries) == 0:
            print("No match for %s" % name)
            print("Names with more than 50 matches cannot be displayed")
            return

        choice = ""
        while not (isinstance(choice, int) and choice >=
                   0 and choice < len(con.entries)):
            choice = get_input("Select person 0-%d (a aborts): " %
                               (len(con.entries) - 1))
            if choice == "a":
                return
            try:
                choice = int(choice)
            except ValueError:
                pass
        entry = con.entries[choice]
        print_person(entry)


if __name__ == '__main__':
    while True:
        print('Søk på navn.')
        print('Trunkeringstegn: *')
        try:
            prompt()
        except KeyboardInterrupt:
            exit(0)
        except Exception as e:
            print(e)
